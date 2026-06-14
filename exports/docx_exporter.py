from docx import Document
import os
import re
from datetime import datetime


def export_docx(report, query):

    os.makedirs(
        "reports",
        exist_ok=True
    )

    timestamp = datetime.now().strftime(
        "%Y%m%d_%H%M%S"
    )

    safe_query = re.sub(
        r"[^a-zA-Z0-9_]",
        "",
        query.replace(" ", "_")
    )

    filepath = (
        f"reports/{safe_query}_{timestamp}.docx"
    )

    doc = Document()

    doc.add_heading(
        "AI Research Report",
        level=1
    )

    doc.add_paragraph(report)

    doc.save(filepath)

    return filepath