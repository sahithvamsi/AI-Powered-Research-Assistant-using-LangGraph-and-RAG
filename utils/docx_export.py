from docx import Document


def create_docx(
    report,
    filename="report.docx"
):

    doc = Document()

    doc.add_heading(
        "Research Report",
        level=1
    )

    doc.add_paragraph(
        report
    )

    doc.save(
        filename
    )

    return filename